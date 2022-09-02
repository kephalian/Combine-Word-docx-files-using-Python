from docxcompose.composer import Composer
from docx import Document
import os, sys
import tempfile
import PySimpleGUI as sg

def resource_path(relative_path):
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


temp_file=tempfile.NamedTemporaryFile(prefix="Temp_master")
temp_file_name=temp_file.name
temp_file.close()
sg.theme("material2")
f_name=sg.popup_get_text("Input file name to save as Joined/ Merged DOCX", title="Input merged DOCX file name",default_text="result")
if f_name is None: sys.exit()
if len(f_name)<=0: sys.exit()
f_name=f_name+".docx"
files = [a for a in os.listdir() if a.endswith(".docx") or a.endswith(".DOCX")]
if os.path.isfile(temp_file_name) is False:
        document2 = Document(resource_path("default.docx"))
        #document2.add_heading('Joined the docx', 0)
        p2 = document2.add_paragraph('\n\r')
        document2.add_page_break()
        document2.save(temp_file_name)
        
master = Document(temp_file_name)
composer = Composer(master)
for file_ in files:
    if file_==temp_file_name: continue
    doc1 = Document(file_)
    composer.append(doc1)


composer.save(f_name)
#sg.popup(temp_file_name)
os.remove(temp_file_name)
sg.popup(f"Saved {f_name}")
